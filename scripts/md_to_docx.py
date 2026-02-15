import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_file, docx_file):
    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        if line.startswith('# '):
            # Título 1
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # Título 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Título 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            # Lista
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # Negrita párrafo completo (simplificado)
            p = doc.add_paragraph()
            runner = p.add_run(line.replace('**', ''))
            runner.bold = True
        elif line.startswith('|'):
            # Tablas (muy básico, solo texto plano para no romper)
            doc.add_paragraph(line, style='No Spacing')
        else:
            # Párrafo normal
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Generado: {docx_file}")

if __name__ == "__main__":
    base_path = "/root/Ticketera"
    files = [
        ("ANEXO_VIII_RESULTADOS_TESTING.md", "ANEXO_VIII_RESULTADOS_TESTING.docx"),
        ("IF_ELEVACION_PROYECTO_TICKETERA.md", "IF_ELEVACION_PROYECTO_TICKETERA.docx")
    ]
    
    for md, docx in files:
        md_path = os.path.join(base_path, md)
        docx_path = os.path.join(base_path, docx)
        if os.path.exists(md_path):
            markdown_to_docx(md_path, docx_path)
        else:
            print(f"No encontrado: {md_path}")
