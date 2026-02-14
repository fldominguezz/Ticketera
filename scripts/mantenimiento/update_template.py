from docx import Document
from docx.shared import Pt
import os

def update_template():
    path = "templates/parte_informativo_base.docx"
    if not os.path.exists(path):
        print(f"Error: {path} not found")
        return

    doc = Document(path)

    # 1. Update Header
    for p in doc.paragraphs:
        if "FECHA:" in p.text:
            p.text = "FECHA: {{FECHA_LARGA}} – {{TURNO}}"
            break

    # 2. Update Bullets
    replacements = {
        "La totalidad de las licencias: ?/700": "La totalidad de las licencias: {{ESET_SOC_LIC_USADAS}}/700",
        "La totalidad de las licencias Mobile: ?/700": "La totalidad de las licencias Mobile: {{ESET_SOC_MOBILE_USADAS}}/700",
        "La totalidad de las licencias: ?/1550": "La totalidad de las licencias: {{ESET_BIENESTAR_LIC_USADAS}}/1550",
        "La totalidad de los incidentes: -?-": "La totalidad de los incidentes: {{ESET_BIENESTAR_INCIDENTES}}",
        "Colectores en Workstation: ?": "Colectores en Workstation: {{EDR_COLECTORES_WS}}",
        "Colectores en Servidores: ?": "Colectores en Servidores: {{EDR_COLECTORES_SRV}}",
        "Licencias desplegadas: ?/1000": "Licencias desplegadas: {{EMS_LIC_USADAS}}/1000",
        "División SEGURIDAD EN REDES DE DATOS: -?-": "División SEGURIDAD EN REDES DE DATOS: {{BLOQUEO_SRD}}",
        "División CENTRO FEDERAL DE DATOS: -?-": "División CENTRO FEDERAL DE DATOS: {{BLOQUEO_CFD}}"
    }

    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    # 3. Update Sections
    # We will clear the existing sections from "NOVEDADES GENERALES" onwards and rebuild them to be sure
    
    # Find start index of sections
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "NOVEDADES GENERALES" in p.text:
            start_idx = i
            break
    
    if start_idx != -1:
        # Delete paragraphs from start_idx to the end
        # doc.paragraphs is a list, but deleting from it is tricky.
        # Better approach: find and replace existing ones, append missing ones.
        pass

    # Actually, the user wants me to be careful with styles.
    # Let's map existing titles to placeholders. 
    
    sections_map = {
        "NOVEDADES GENERALES": "{{NOVEDADES_GENERALES}}",
        "ESET CLOUD": "{{ESET_SOC_OBS}}\nEstado de salud: {{ESET_SOC_SALUD}}",
        "ESET BIENESTAR": "{{ESET_BIENESTAR_OBS}}\nEstado de salud: {{ESET_BIENESTAR_SALUD}}",
        "FORTISANDBOX": "{{FORTISANDBOX_OBS}}\nEstado de salud: {{FORTISANDBOX_SALUD}}",
        "FORTIEMS": "{{EMS_OBS}}\nEstado de salud: {{EMS_SALUD}}"
    }

    # Track which ones we found
    found_sections = set()

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in sections_map:
            found_sections.add(text)
            # The next paragraph (or the one after) usually has the "?"
            # Let's find the "?" in the following paragraphs until the next title or empty space
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                if "?" in doc.paragraphs[j].text:
                    doc.paragraphs[j].text = sections_map[text]
                    break
    
    # Now add missing sections
    all_required_sections = [
        ("6) FORTISIEM", "{{FORTISIEM_OBS}}\nEstado de salud: {{FORTISIEM_SALUD}}"),
        ("7) FORTIANALYZER", "{{FORTIANALYZER_OBS}}\nEstado de salud: {{FORTIANALYZER_SALUD}}"),
        ("8) FORTIEDR", "{{FORTIEDR_OBS}}\nEstado de salud: {{FORTIEDR_SALUD}}"),
        ("9) CORREO POLICIAL", "{{CORREO_OBS}}")
    ]

    # Style from FORTIEMS (para 42/43)
    # We will try to copy style if possible, but python-docx is limited in "copying styles" easily between paras.
    # We can at least use similar formatting.

    for title, content in all_required_sections:
        # Check if already exists (maybe with different numbering)
        exists = False
        clean_title = title.split(")")[-1].strip()
        for p in doc.paragraphs:
            if clean_title in p.text:
                exists = True
                break
        
        if not exists:
            doc.add_paragraph("") # Spacer
            p_title = doc.add_paragraph(title)
            p_title.style = doc.paragraphs[42].style # Try to match FORTIEMS style
            p_content = doc.add_paragraph(content)

    doc.save(path)
    print("Template updated successfully")

if __name__ == "__main__":
    update_template()
