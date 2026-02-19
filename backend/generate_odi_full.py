from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    t = doc.add_heading(text, level=level)
    if level == 1:
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run = t.runs[0]
        run.font.color.rgb = RGBColor(0, 51, 102)

def generate_exhaustive_document():
    doc = Document()
    
    # --- PORTADA ---
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POLICÍA FEDERAL ARGENTINA")
    run.bold = True
    run.font.size = Pt(22)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUPERINTENDENCIA FEDERAL DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIONES")
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DIRECCIÓN GENERAL DE TECNOLOGÍAS DE LA INFORMACIÓN")
    run.bold = True
    run.font.size = Pt(12)
    
    for _ in range(3): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EXPEDIENTE TÉCNICO INTEGRAL DE SOFTWARE")
    run.bold = True
    run.font.size = Pt(26)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTEMA: TICKETERA SOC v1.0")
    run.font.size = Pt(20)
    
    for _ in range(5): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFORME A DIRECTIVA P001 - ODI Nº 058/2024")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_page_break()

    # --- INDICE ---
    add_title(doc, "ÍNDICE ESTRUCTURAL", level=1)
    topics = [
        "1. INTRODUCCIÓN", "2. OBJETIVO", "3. ALCANCE", "4. VIGENCIA",
        "5. LINEAMIENTOS GENERALES", "6. CALIDAD", "7. IMPLEMENTACIÓN", 
        "8. DOCUMENTACIÓN", "9. BACKUP", "10. POST-IMPLEMENTACIÓN", 
        "11. DISPOSICIÓN FINAL Y ARCHIVO", "12. LICENCIAS", "13. GLOSARIO", 
        "14. ANEXO I - SOLICITUD DE PROYECTO"
    ]
    for t in topics:
        doc.add_paragraph(t)
    doc.add_page_break()

    # --- CONTENIDO ---
    add_title(doc, "1. INTRODUCCIÓN", level=2)
    doc.add_paragraph("La presente documentación técnica detalla el desarrollo y arquitectura del sistema 'Ticketera SOC', diseñado para la División Seguridad Informática de la Policía Federal Argentina. El sistema cumple íntegramente con las directivas establecidas en la ODI 058/24, garantizando la estandarización institucional en la creación de software.")

    add_title(doc, "2. OBJETIVO", level=2)
    doc.add_paragraph("El objetivo primordial es dotar a la institución de una plataforma centralizada y soberana para la gestión de incidentes de ciberseguridad. Se busca optimizar los tiempos de respuesta (SLA) y asegurar la integridad de las evidencias recolectadas en cada ticket.")

    add_title(doc, "3. ALCANCE", level=2)
    doc.add_paragraph("Ticketera SOC alcanza operativamente a todos los analistas del Centro de Operaciones de Seguridad (SOC) y administrativamente a la Dirección de Ciberseguridad.")

    add_title(doc, "5. LINEAMIENTOS GENERALES", level=2)
    doc.add_paragraph("5.1. PROPIEDAD: El software es propiedad exclusiva de la PFA (Cap. 5.1.3). El código fuente se encuentra auditado y alojado en la infraestructura del CENTRO FEDERAL DE DATOS.")
    
    add_title(doc, "5.4. Metodología del Ciclo de Vida", level=3)
    doc.add_paragraph("Se ha implementado el marco conceptual 'Ciclo de vida del desarrollo de Software' exigido por la normativa:")
    doc.add_paragraph("I. RELEVAMIENTO Y ANÁLISIS: Definición de flujos de trabajo con el personal técnico del SOC.", style='List Bullet')
    doc.add_paragraph("II. DISEÑO Y DESARROLLO: Programación modular bajo arquitectura de microservicios.", style='List Bullet')
    doc.add_paragraph("III. PRUEBA / PREPRODUCCIÓN: Testing exhaustivo en ambiente controlado.", style='List Bullet')
    doc.add_paragraph("IV. IMPLEMENTACIÓN: Despliegue automatizado mediante contenedores Docker.", style='List Bullet')

    add_title(doc, "5.5. Seguridad de la Información", level=3)
    doc.add_paragraph("Cumplimiento total con la Política de Seguridad Institucional. Medidas aplicadas:")
    doc.add_paragraph("• CIFRADO: Almacenamiento de credenciales mediante Argon2id con sal aleatoria.", style='List Bullet')
    doc.add_paragraph("• AUTENTICACIÓN: Multi-Factor Authentication (MFA) vía TOTP obligatorio para analistas.", style='List Bullet')
    doc.add_paragraph("• INTEGRIDAD: Verificación humana local para prevenir ataques automatizados.", style='List Bullet')
    doc.add_paragraph("• SESIONES: Gestión de tokens JWT con tiempo de vida limitado y revocación dinámica.", style='List Bullet')

    add_title(doc, "5.10. Modularidad", level=3)
    doc.add_paragraph("El sistema está dividido en módulos independientes conforme a la Cláusula 5.10.1:")
    doc.add_paragraph("• Módulo de Gestión de Tickets: Núcleo de la aplicación (FastAPI).", style='List Bullet')
    doc.add_paragraph("• Módulo de Integración SIEM: Receptor de alertas externas (SOC-Module).", style='List Bullet')
    doc.add_paragraph("• Módulo Forense: Análisis de archivos EML y hashes (VirusTotal API).", style='List Bullet')
    doc.add_paragraph("• Interfaz de Usuario: SPA construida en React con Next.js.", style='List Bullet')

    add_title(doc, "8. DOCUMENTACIÓN - DICCIONARIO DE DATOS", level=2)
    doc.add_paragraph("Conforme a la Cláusula 8.1.a, se describe la estructura de datos crítica del sistema:")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Entidad'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Nulidad'
    hdr[3].text = 'Definición'
    
    fields = [
        ('tickets.id', 'UUID', 'No', 'Clave primaria inmutable.'),
        ('tickets.status', 'Enum', 'No', 'Estado operativo del incidente.'),
        ('users.hashed_pw', 'String', 'No', 'Hash de seguridad del operador.'),
        ('audit_logs.ip', 'String', 'No', 'IP de origen para auditoría forense.'),
        ('assets.mac', 'String', 'Sí', 'Identificación física del equipo.')
    ]
    for f, t, n, d in fields:
        row = table.add_row().cells
        row[0].text = f
        row[1].text = t
        row[2].text = n
        row[3].text = d

    add_title(doc, "13. GLOSARIO TÉCNICO", level=2)
    glossary = [
        ("Ambiente", "Entorno tecnológico donde opera la solución."),
        ("Backup", "Copia de respaldo cifrada de la información institucional."),
        ("Encriptación", "Proceso de protección de datos mediante algoritmos matemáticos."),
        ("Repositorio", "Almacenamiento de versiones del código fuente (Git)."),
        ("SLA", "Acuerdo de Niveles de Servicio para la resolución de tickets.")
    ]
    for term, definition in glossary:
        doc.add_paragraph(f"{term}: {definition}", style='List Bullet')

    add_title(doc, "14. ANEXO I - SOLICITUD DE PROYECTO", level=2)
    doc.add_paragraph("PROYECTO: Ticketera SOC v1.0")
    doc.add_paragraph("DEPENDENCIA SOLICITANTE: División Seguridad Informática")
    add_title(doc, "Resumen Ejecutivo", level=3)
    doc.add_paragraph("Implementación de un gestor de incidentes basado en inteligencia artificial para la Policía Federal Argentina. El sistema centraliza alertas de SIEM, permite el triaje automatizado y garantiza la trazabilidad total de las acciones de los analistas de ciberseguridad.")

    # Guardar
    output_path = "/tmp/EXPEDIENTE_TECNICO_ODI_058_EXHAUSTIVO.docx"
    doc.save(output_path)

if __name__ == "__main__":
    generate_exhaustive_document()
