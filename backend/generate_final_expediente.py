from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    t = doc.add_heading(text, level=level)
    if level == 1:
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 51, 102)

def add_par(doc, text, bold=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def generate():
    doc = Document()
    
    # --- PORTADA ---
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POLICÍA FEDERAL ARGENTINA")
    run.bold = True
    run.font.size = Pt(22)
    
    add_par(doc, "SUPERINTENDENCIA FEDERAL DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIONES").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_par(doc, "DIRECCIÓN GENERAL DE TECNOLOGÍAS DE LA INFORMACIÓN").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(4): doc.add_paragraph()
    add_par(doc, "EXPEDIENTE TÉCNICO DE CUMPLIMIENTO INTEGRAL").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_par(doc, "PROYECTO: TICKETERA SOC v1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5): doc.add_paragraph()
    add_par(doc, "CONFORME A DIRECTIVA P001 - ODI Nº 058/2024").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- ÍNDICE ---
    add_title(doc, "ÍNDICE", level=1)
    items = [
        "1. INTRODUCCIÓN", "2. OBJETIVO", "3. ALCANCE", "4. VIGENCIA",
        "5. LINEAMIENTOS GENERALES", "6. CALIDAD", "7. IMPLEMENTACIÓN",
        "8. DOCUMENTACIÓN", "9. BACKUP", "10. POST-IMPLEMENTACIÓN",
        "11. DISPOSICIÓN FINAL Y ARCHIVO", "12. LICENCIAS", "13. GLOSARIO",
        "14. ANEXO I - SOLICITUD DE PROYECTO"
    ]
    for i in items:
        add_par(doc, i)
    doc.add_page_break()

    # --- DESARROLLO ---
    add_title(doc, "1. INTRODUCCIÓN", level=2)
    add_par(doc, "La Superintendencia FEDERAL DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIONES establece mediante la ODI 058/24 el marco de referencia obligatorio para el desarrollo de software institucional. El sistema 'Ticketera SOC' se alinea con estos estándares para proporcionar una herramienta centralizada y segura a la División Seguridad Informática.")

    add_title(doc, "2. OBJETIVO", level=2)
    add_par(doc, "Establecer la base técnica y normativa de Ticketera SOC para:")
    add_par(doc, "2.1. Disponer de una metodología homogénea para las actividades de respuesta ante incidentes.", style='List Bullet')
    add_par(doc, "2.2. Estandarizar el ciclo de vida del software institucional.", style='List Bullet')
    add_par(doc, "2.3. Garantizar la integridad de la información policial mediante procesos auditables.", style='List Bullet')

    add_title(doc, "5. LINEAMIENTOS GENERALES", level=2)
    add_par(doc, "5.1.1. Adaptación a la estructura orgánica institucional.")
    add_par(doc, "5.1.3. Propiedad intelectual de la Policía Federal Argentina.")
    add_par(doc, "5.1.5. Código fuente alojado en el repositorio Git institucional.")
    
    add_title(doc, "5.4. Metodología", level=3)
    add_par(doc, "Se ha cumplido el Ciclo de Vida: I. Relevamiento, II. Diseño, III. Prueba, IV. Implementación, V. Soporte, VI. Archivo.")

    add_title(doc, "5.5. Seguridad", level=3)
    add_par(doc, "5.5.1. Cumplimiento con la Política de Seguridad Institucional.")
    add_par(doc, "5.5.2. Privacidad de datos según normativa DNPDP. Implementación de MFA y Hashing Argon2.")

    add_title(doc, "8. DOCUMENTACIÓN - DICCIONARIO DE DATOS", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Entidad'
    hdr[1].text = 'Columna'
    hdr[2].text = 'Tipo'
    hdr[3].text = 'Descripción'
    
    db_rows = [
        ('users', 'id', 'UUID', 'Clave primaria.'),
        ('users', 'username', 'VARCHAR', 'Identificador.'),
        ('tickets', 'status', 'VARCHAR', 'Estado operativo.'),
        ('audit_logs', 'event_type', 'VARCHAR', 'Tipo de evento.'),
        ('assets', 'ip_address', 'VARCHAR', 'Dirección IP.')
    ]
    for r in db_rows:
        row = table.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v

    add_title(doc, "13. GLOSARIO", level=2)
    terms = [
        ("Ambiente", "Entorno tecnológico de la aplicación."),
        ("Backup", "Copia de respaldo cifrada."),
        ("Encriptación", "Protección mediante algoritmos."),
        ("Repositorio", "Almacenamiento de versiones (Git).")
    ]
    for t, d in terms:
        add_par(doc, f"{t}: {d}", style='List Bullet')

    add_title(doc, "14. ANEXO I - SOLICITUD DE PROYECTO", level=1)
    add_par(doc, "RESUMEN EJECUTIVO:", bold=True)
    add_par(doc, "Ticketera SOC es un desarrollo soberano que centraliza la gestión de incidentes, automatiza el triaje de alertas de SIEM y garantiza auditoría total conforme a la ODI 058/24.")

    doc.save("/tmp/EXPEDIENTE_TECNICO_ODI_058_FINAL.docx")

if __name__ == "__main__":
    generate()
