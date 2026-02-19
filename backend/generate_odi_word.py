from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def add_title(doc, text, level=1):
    t = doc.add_heading(text, level=level)
    if level == 1:
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

def generate_document():
    doc = Document()
    
    # --- PORTADA ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POLICÍA FEDERAL ARGENTINA")
    run.bold = True
    run.font.size = Pt(16)
    p.add_run("\n")
    
    run = p.add_run("DIRECCIÓN GENERAL DE TECNOLOGÍAS DE LA INFORMACIÓN")
    run.bold = True
    run.font.size = Pt(14)
    p.add_run("\n\n")
    
    run = p.add_run("EXPEDIENTE TÉCNICO DE PROYECTO DE SOFTWARE")
    run.bold = True
    run.font.size = Pt(20)
    p.add_run("\n")
    
    run = p.add_run("SISTEMA: TICKETERA SOC v1.0")
    run.font.size = Pt(18)
    p.add_run("\n\n")
    
    run = p.add_run("CONFORME A ODI 058/24 - DIRECTIVA P001")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_page_break()

    # --- INDICE ---
    add_title(doc, "ÍNDICE", level=1)
    indice = [
        "1. INTRODUCCIÓN", "2. OBJETIVO", "3. ALCANCE", "4. VIGENCIA", 
        "5. LINEAMIENTOS GENERALES", "6. CALIDAD", "7. IMPLEMENTACIÓN", 
        "8. DOCUMENTACIÓN", "9. BACKUP", "10. POST-IMPLEMENTACIÓN", 
        "11. DISPOSICIÓN FINAL Y ARCHIVO", "12. LICENCIAS", "13. GLOSARIO", 
        "14. ANEXO I - SOLICITUD DE PROYECTO"
    ]
    for item in indice:
        doc.add_paragraph(item)
    doc.add_page_break()

    # --- CAPITULO 1 AL 4 ---
    add_title(doc, "1. INTRODUCCIÓN", level=2)
    doc.add_paragraph("La Superintendencia FEDERAL DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIONES establece mediante la ODI 058/24 el marco normativo para el desarrollo de aplicativos. El sistema Ticketera SOC se presenta como la solución integral para la gestión de incidentes de ciberseguridad.")

    add_title(doc, "2. OBJETIVO", level=2)
    doc.add_paragraph("Definir lineamientos y mejores prácticas para el desarrollo del software con el fin de disponer de una metodología homogénea, estandarizar procesos y garantizar la calidad y seguridad institucional.")

    add_title(doc, "3. ALCANCE", level=2)
    doc.add_paragraph("El desarrollo alcanza a la División SEGURIDAD INFORMÁTICA, proporcionando herramientas de triaje, análisis forense y seguimiento de tickets de seguridad.")

    add_title(doc, "4. VIGENCIA", level=2)
    doc.add_paragraph("A partir de su aprobación y puesta en marcha en febrero de 2026.")

    # --- CAPITULO 5: LINEAMIENTOS GENERALES ---
    add_title(doc, "5. LINEAMIENTOS GENERALES", level=2)
    doc.add_paragraph("5.1.1. Se ha considerado la estructura orgánica institucional afectando a la División Seguridad Informática.")
    doc.add_paragraph("5.1.3. Los componentes son propiedad de la Policía Federal Argentina, quedando los derechos de autor para la Institución.")
    doc.add_paragraph("5.1.5. El código fuente reside en la División CENTRO FEDERAL DE DATOS.")
    
    add_title(doc, "5.4. Metodología", level=3)
    doc.add_paragraph("Se ha seguido el Ciclo de Vida del Software: I. Relevamiento, II. Diseño y Desarrollo, III. Prueba (Preproducción), IV. Implementación, V. Post-Implementación, VI. Disposición final.")

    add_title(doc, "5.5. Seguridad", level=3)
    doc.add_paragraph("Se cumple con la 'Política de Seguridad de la Información' vigente. Implementación de hashing Argon2, 2FA y auditoría inmutable.")

    add_title(doc, "5.10. Modularidad", level=3)
    doc.add_paragraph("La programación es modular, dividida en microservicios independientes (Backend FastAPI, Frontend React, SOC Module Python) que interactúan a través de APIs documentadas.")

    # --- CAPITULO 6: CALIDAD ---
    add_title(doc, "6. CALIDAD", level=2)
    doc.add_paragraph("Se dispone de un Ambiente y Plan de Pruebas. Los registros de ejecución de pruebas unitarias y de integración han sido validados con éxito en el entorno de preproducción.")

    # --- CAPITULO 8: DOCUMENTACIÓN ---
    add_title(doc, "8. DOCUMENTACIÓN", level=2)
    doc.add_paragraph("Se entregan los manuales técnicos, de instalación y de usuario conformados previamente por el área usuaria.")
    
    add_title(doc, "8.1.a. Diccionario de Datos (Resumen de Tablas Críticas)", level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tabla'
    hdr_cells[1].text = 'Función'
    hdr_cells[2].text = 'Datos Sensibles'
    
    data_tables = [
        ('users', 'Gestión de analistas y permisos', 'Emails, Hashed PW'),
        ('tickets', 'Registro de incidentes de seguridad', 'Contenido del incidente'),
        ('assets', 'Inventario de infraestructura crítica', 'IPs, MACs, Hostnames'),
        ('audit_logs', 'Registro inmutable de acciones', 'Actividad de usuario'),
        ('siem_alerts', 'Alertas provenientes de FortiSIEM', 'Logs crudos de red')
    ]
    
    for tab, func, sens in data_tables:
        row_cells = table.add_row().cells
        row_cells[0].text = tab
        row_cells[1].text = func
        row_cells[2].text = sens

    # --- CAPITULO 9: BACKUP ---
    add_title(doc, "9. BACKUP", level=2)
    doc.add_paragraph("9.3. El backup se realiza en medio físico distinto de la estación de trabajo, garantizando encriptación que evita el robo o alteración de datos.")

    # --- CAPITULO 13: GLOSARIO ---
    add_title(doc, "13. GLOSARIO", level=2)
    glosario = [
        ("Ambiente", "Entorno de condiciones en el cual opera la aplicación."),
        ("Backup", "Copia de seguridad realizada sobre ficheros o bases de datos."),
        ("Encriptación", "Proceso de convertir información en un formato ilegible mediante algoritmos."),
        ("Repositorio", "Lugar centralizado donde se almacenan y gestionan versiones de código fuente."),
        ("Standalone", "Aplicación que funciona de manera independiente.")
    ]
    for term, desc in glosario:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{term}: ").bold = True
        p.add_run(desc)

    # --- ANEXO I ---
    doc.add_page_break()
    add_title(doc, "14. ANEXO I - SOLICITUD DE PROYECTO", level=1)
    
    add_title(doc, "Resumen Ejecutivo", level=3)
    doc.add_paragraph("Implementación de una plataforma soberana para la gestión de incidentes, automatización de triaje y análisis de amenazas asistido por IA.")
    
    add_title(doc, "Alcance del Requerimiento", level=3)
    doc.add_paragraph("Gestión de tickets, integración SIEM, laboratorio forense, wiki de procedimientos y dashboard gerencial.")

    # Guardar
    output_path = "/tmp/EXPEDIENTE_TECNICO_ODI_058.docx"
    doc.save(output_path)

if __name__ == "__main__":
    generate_document()
