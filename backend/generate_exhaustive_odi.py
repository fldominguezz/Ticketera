from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    t = doc.add_heading(text, level=level)
    if level == 1:
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 51, 102)

def add_par(doc, text, bold=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def generate():
    doc = Document()
    
    # --- PORTADA ---
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POLICÍA FEDERAL ARGENTINA")
    run.bold = True
    run.font.size = Pt(22)
    
    add_par(doc, "SUPERINTENDENCIA FEDERAL DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIONES").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_par(doc, "DIRECCIÓN GENERAL DE TECNOLOGÍAS DE LA INFORMACIÓN").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(4): doc.add_paragraph()
    add_par(doc, "EXPEDIENTE TÉCNICO DE CUMPLIMIENTO INTEGRAL").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_par(doc, "PROYECTO: TICKETERA SOC v1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5): doc.add_paragraph()
    add_par(doc, "CONFORME A DIRECTIVA P001 - ODI Nº 058/2024").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- CONTENIDO DETALLADO PUNTO POR PUNTO ---
    
    add_title(doc, "1. INTRODUCCIÓN", level=2)
    add_par(doc, "La Superintendencia FEDERAL DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIONES establece el marco normativo para el desarrollo de aplicativos. Ticketera SOC se alinea con la estrategia de ciberdefensa nacional.")

    add_title(doc, "2. OBJETIVO", level=2)
    add_par(doc, "Definir lineamientos técnicos bajo la ODI 058/24 para asegurar:")
    add_par(doc, "2.1. Metodología homogénea.", style='List Bullet')
    add_par(doc, "2.2. Estandarización de procesos SOC.", style='List Bullet')
    add_par(doc, "2.3. Calidad y seguridad en la gestión de incidentes.", style='List Bullet')

    add_title(doc, "5. LINEAMIENTOS GENERALES", level=2)
    add_par(doc, "5.1.1. Adaptación a la estructura orgánica de la División Seguridad Informática.")
    add_par(doc, "5.1.3. La propiedad intelectual pertenece a la Policía Federal Argentina.")
    add_par(doc, "5.1.5. El código fuente reside en la División CENTRO FEDERAL DE DATOS.")
    
    add_title(doc, "5.4. Metodología", level=3)
    add_par(doc, "Se ha cumplido el Ciclo de Vida: I. Relevamiento, II. Diseño, III. Prueba, IV. Implementación, V. Soporte, VI. Archivo.")

    add_title(doc, "5.5. Seguridad", level=3)
    add_par(doc, "5.5.1. Cumplimiento con la Política de Seguridad de la Información.")
    add_par(doc, "5.5.2. Privacidad de datos según Disposición DNPDP Nº 18/2015. Implementación de MFA y Hashing Argon2.")

    add_title(doc, "5.10. Modularidad", level=3)
    add_par(doc, "Programación modular dividida en microservicios: API REST (FastAPI), Interfaz (React), Integrador SIEM.")

    add_title(doc, "8. DOCUMENTACIÓN - DICCIONARIO DE DATOS", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Entidad'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Nulidad'
    hdr[3].text = 'Descripción'
    
    data = [
        ('tickets', 'UUID', 'No', 'Clave única del incidente.'),
        ('status', 'Enum', 'No', 'Estado operativo.'),
        ('users', 'UUID', 'No', 'ID del analista.'),
        ('audit_logs', 'JSONB', 'No', 'Registro inmutable.'),
        ('assets', 'String', 'Sí', 'IP/MAC del equipo.')
    ]
    for e, t, n, d in data:
        row = table.add_row().cells
        row[0].text = e
        row[1].text = t
        row[2].text = n
        row[3].text = d

    add_title(doc, "13. GLOSARIO", level=2)
    terms = [
        ("Ambiente", "Entorno tecnológico de la aplicación."),
        ("Backup", "Copia de seguridad cifrada."),
        ("Encriptación", "Protección de datos mediante algoritmos."),
        ("Repositorio", "Almacenamiento de código Git."),
        ("SLA", "Acuerdos de Nivel de Servicio.")
    ]
    for t, d in terms:
        add_par(doc, f"{t}: {d}", style='List Bullet')

    add_title(doc, "14. ANEXO I - SOLICITUD DE PROYECTO", level=1)
    add_par(doc, "RESUMEN EJECUTIVO:", bold=True)
    add_par(doc, "Ticketera SOC centraliza la gestión de incidentes de la PFA, automatizando el triaje de alertas de SIEM y garantizando auditoría total.")

    doc.save("/tmp/EXPEDIENTE_TECNICO_MAESTRO_ODI_058.docx")

if __name__ == "__main__":
    generate()
