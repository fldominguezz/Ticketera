import os
import shutil
import re
from datetime import date
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List

# Constants for License Maxs
LICENSE_MAXS = {
    "ESET_SOC_LIC_MAX": 700,
    "ESET_SOC_MOBILE_LIC_MAX": 700,
    "ESET_BIENESTAR_LIC_MAX": 1550,
    "EMS_LIC_MAX": 1000
}

class DailyReportGenerator:
    def __init__(self, template_path: str = "templates/parte_informativo_base.docx"):
        self.template_path = template_path

    def _get_long_date(self, d: date) -> str:
        days = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]
        months = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
        day_str = days[d.weekday()]
        month_str = months[d.month - 1]
        return f"{day_str} {d.day} de {month_str} del {d.year}"

    def _copy_para_format(self, src_p, dst_p):
        """Copies paragraph formatting from src_p to dst_p strictly."""
        if not src_p or not dst_p: return
        try:
            dst_p.style = src_p.style
        except:
            pass
        dst_fmt = dst_p.paragraph_format
        src_fmt = src_p.paragraph_format
        
        # Align left to avoid "huge spaces" from justification
        dst_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        dst_fmt.first_line_indent = src_fmt.first_line_indent
        dst_fmt.left_indent = src_fmt.left_indent
        dst_fmt.right_indent = src_fmt.right_indent
        dst_fmt.space_after = src_fmt.space_after
        dst_fmt.space_before = src_fmt.space_before
        dst_fmt.line_spacing = src_fmt.line_spacing

    def _set_para_text_clean(self, paragraph, text, model_p=None):
        """Replaces text in a paragraph run by run, sanitizing and removing tab stops."""
        # 1. Sanitize text: collapse tabs and multiple spaces to ONE space
        clean_text = re.sub(r'\s+', ' ', text.replace('\t', ' ').replace('\r', '')).strip()
        
        # 2. Force Left Alignment to prevent justification issues
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 3. Remove all tab stops via XML to prevent column behavior
        pPr = paragraph._p.get_or_add_pPr()
        for tabs in pPr.xpath('.//w:tabs'):
            pPr.remove(tabs)
            
        # 4. Replace text in runs
        # Try to find a run with a placeholder, otherwise clear and add one
        found = False
        placeholder_patterns = [r'\?', r'\{\{.*\}\}']
        
        for run in paragraph.runs:
            for pattern in placeholder_patterns:
                if re.search(pattern, run.text):
                    run.text = clean_text
                    found = True
                    break
            if found: break
            
        if not found:
            # If no placeholder run found, clear all and add a new one with model font
            p_element = paragraph._p
            for run_element in p_element.xpath('.//w:r'):
                p_element.remove(run_element)
            
            new_run = paragraph.add_run(clean_text)
            if model_p and model_p.runs:
                src_run = model_p.runs[0]
                new_run.bold = src_run.bold
                new_run.italic = src_run.italic
                new_run.underline = src_run.underline
                if src_run.font.name: new_run.font.name = src_run.font.name
                if src_run.font.size: new_run.font.size = src_run.font.size

    def _insert_paragraph_after(self, paragraph):
        """Insert a new empty paragraph after the given paragraph."""
        new_p_xml = OxmlElement('w:p')
        paragraph._p.addnext(new_p_xml)
        return paragraph.__class__(new_p_xml, paragraph._parent)

    def _ensure_sections_exist(self, doc_path: str):
        doc = Document(doc_path)
        all_paras = list(doc.paragraphs)
        
        detailed_start_idx = 0
        for i, p in enumerate(all_paras):
            if "NOVEDADES GENERALES" in p.text.upper():
                detailed_start_idx = i
                break
        
        ref_title_p = None
        ref_content_p = None
        
        for i in range(detailed_start_idx, len(all_paras)):
            p = all_paras[i]
            t = p.text.upper()
            if any(k in t for k in ["FORTISANDBOX", "FORTIEMS", "NOVEDADES GENERALES"]):
                ref_title_p = p
                for j in range(i + 1, len(all_paras)):
                    if all_paras[j].text.strip():
                        ref_content_p = all_paras[j]
                        break
                if ref_content_p: break
        pass
        if not ref_title_p: ref_title_p = all_paras[detailed_start_idx] if detailed_start_idx < len(all_paras) else all_paras[0]
        if not ref_content_p: ref_content_p = all_paras[detailed_start_idx+1] if detailed_start_idx+1 < len(all_paras) else all_paras[0]

        sections = [
            ("NOVEDADES GENERALES", ["{{NOVEDADES_GENERALES}}"]),
            ("ESET CLOUD", ["{{ESET_SOC_OBS}}", "Estado de salud: {{ESET_SOC_SALUD}}"]),
            ("ESET BIENESTAR", ["{{ESET_BIENESTAR_OBS}}", "Estado de salud: {{ESET_BIENESTAR_SALUD}}"]),
            ("FORTISANDBOX", ["{{FORTISANDBOX_OBS}}", "Estado de salud: {{FORTISANDBOX_SALUD}}"]),
            ("FORTIEMS", ["{{EMS_OBS}}", "Estado de salud: {{EMS_SALUD}}"]),
            ("FORTISIEM", ["{{FORTISIEM_OBS}}", "Estado de salud: {{FORTISIEM_SALUD}}"]),
            ("FORTIANALYZER", ["{{FORTIANALYZER_OBS}}", "Estado de salud: {{FORTIANALYZER_SALUD}}"]),
            ("FORTIEDR", ["{{FORTIEDR_OBS}}", "Estado de salud: {{FORTIEDR_SALUD}}"]),
            ("CORREO POLICIAL", ["{{CORREO_OBS}}"])
        ]

        for keyword, lines in sections:
            found_title_p = None
            current_paras = list(doc.paragraphs)
            for i in range(detailed_start_idx, len(current_paras)):
                p = current_paras[i]
                if keyword in p.text.upper() and len(p.text.strip()) < 45:
                    found_title_p = p
                    break
            
            if found_title_p:
                current_p = found_title_p
                for line_text in lines:
                    next_p = None
                    is_next_title = False
                    sibling = current_p._p.getnext()
                    while sibling is not None:
                        if sibling.tag.endswith('p'):
                            temp_p = current_p.__class__(sibling, current_p._parent)
                            if temp_p.text.strip():
                                next_p = temp_p
                                for k_check, _ in sections:
                                    if k_check in next_p.text.upper() and len(next_p.text.strip()) < 45:
                                        is_next_title = True
                                        break
                                break
                        sibling = sibling.getnext()
                    
                    if not next_p or is_next_title or ("{{" not in next_p.text and "Estado de salud" not in next_p.text and "?" not in next_p.text):
                        new_p = self._insert_paragraph_after(current_p)
                        self._copy_para_format(ref_content_p, new_p)
                        self._set_para_text_clean(new_p, line_text, ref_content_p)
                        current_p = new_p
                    else:
                        self._set_para_text_clean(next_p, line_text, ref_content_p)
                        self._copy_para_format(ref_content_p, next_p)
                        current_p = next_p
            else:
                doc.add_paragraph("")
                p_title = doc.add_paragraph(keyword)
                self._copy_para_format(ref_title_p, p_title)
                for line_text in lines:
                    p_content = doc.add_paragraph(line_text)
                    self._copy_para_format(ref_content_p, p_content)
        
        doc.save(doc_path)

    def generate(self, data: Dict[str, Any], output_path: str) -> str:
        if not os.path.exists(self.template_path):
             alt_path = os.path.join(os.path.dirname(__file__), "../../templates/parte_informativo_base.docx")
             self.template_path = alt_path if os.path.exists(alt_path) else self.template_path
        
        temp_template = output_path + ".temp.docx"
        shutil.copy2(self.template_path, temp_template)
        self._ensure_sections_exist(temp_template)
        
        doc = DocxTemplate(temp_template)
        context = data.copy()
        if 'FECHA_LARGA' not in context and 'date_obj' in context:
             context['FECHA_LARGA'] = self._get_long_date(context['date_obj'])
             
        context.update(LICENSE_MAXS)
        if isinstance(context.get('NOVEDADES_GENERALES'), list):
            context['NOVEDADES_GENERALES'] = "\n\n".join([item for item in context['NOVEDADES_GENERALES'] if item]).strip()

        doc.render(context)
        doc.save(output_path)
        if os.path.exists(temp_template): os.remove(temp_template)
        return output_path

    def extract_text(self, file_path: str) -> str:
        if not os.path.exists(file_path): return ""
        doc = Document(file_path)
        return '\n'.join([p.text for p in doc.paragraphs])

report_generator = DailyReportGenerator()
